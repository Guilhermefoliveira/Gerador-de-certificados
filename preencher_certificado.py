import pkg_resources
import sys
import os
from flask import Flask, request, render_template_string, send_file, redirect, url_for
from docx import Document
from datetime import datetime

app = Flask(__name__, static_folder='static')

def substituir_variaveis(paragraphs, context):
    for p in paragraphs:
        for key, value in context.items():
            if key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, str(value))
                        inline[i].text = text

def load_template(template_name):
    # Verifica se o aplicativo está sendo executado como um executável
    if getattr(sys, 'frozen', False):
        template_path = os.path.join(sys._MEIPASS, 'templates', template_name)
        return open(template_path, 'rb')
    else:
        template_path = os.path.join('templates', template_name)
        return pkg_resources.resource_stream(__name__, template_path)

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Obter os dados do formulário
        nome_aluno = request.form.get('nome-aluno', '')
        nome_curso = request.form.get('nome-curso', '')
        cnh = request.form.get('cnh', '')
        categoria = request.form.get('categoria', '')
        cpf = request.form.get('cpf', '')
        rg = request.form.get('rg', '')
        uf = request.form.get('uf', '')
        renach = request.form.get('renach', '')
        codigo_certificado = request.form.get('codigo-certificado', '')
        id_matricula = request.form.get('id-matricula', '')
        data_inicio = request.form.get('data-inicio', '')
        data_fim = request.form.get('data-fim', '')
        folha = request.form.get('folha', '')
        livro = request.form.get('livro', '')
        a = request.form.get('a', '')
        b = request.form.get('b', '')
        c = request.form.get('c', '')
        d = request.form.get('d', '')
        e = request.form.get('e', '')
        f = request.form.get('f', '')
        data_hoje = request.form.get('data-hoje', '')

        # Obter o nome do template escolhido pelo usuário
        template_name = request.form.get('template', '')

        # Carregar o template do certificado escolhido
        template = Document(load_template(template_name))

        # Contexto com os dados a serem substituídos no template
        context = {
            '{nome_aluno}': nome_aluno,
            '{nome_curso}': nome_curso,
            '{cnh}': cnh,
            '{categoria}': categoria,
            '{cpf}': cpf,
            '{rg}': rg,
            '{uf}': uf,
            '{renach}': renach,
            '{codigo_certificado}': codigo_certificado,
            '{id_matricula}': id_matricula,
            '{data_inicio}': data_inicio,
            '{data_fim}': data_fim,
            '{folha}': folha,
            '{livro}': livro,
            '{nota1}': a,
            '{nota2}': b,
            '{nota3}': c,
            '{nota4}': d,
            '{nota5}': e,
            '{nota6}': f,
            '{data_hoje}': data_hoje
        }

        # Substituir as variáveis no documento (parágrafos fora da tabela)
        for p in template.paragraphs:
            substituir_variaveis([p], context)

        # Substituir as variáveis nas tabelas do documento
        for table in template.tables:
            for row in table.rows:
                for cell in row.cells:
                    substituir_variaveis(cell.paragraphs, context)

        # Gerar um nome único para o novo arquivo
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        novo_nome_arquivo = f'certificado_{nome_aluno}.docx'

        # Caminho completo para o novo arquivo
        output_dir = 'Certificados_Gerados'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        output_path = os.path.join(output_dir, novo_nome_arquivo)

        # Salvar o novo arquivo com os dados preenchidos
        try:
            template.save(output_path)
        except Exception as e:
            return f"Erro ao salvar o certificado: {str(e)}"

        # Redirecionar para a página de sucesso
        return redirect(url_for('success', filename=novo_nome_arquivo))

    # Se o método for GET, exibir o formulário
    template = load_template('index.html').read().decode('utf-8')
    return render_template_string(template)

@app.route('/success')
def success():
    # Recupera o nome do arquivo a partir da query string
    filename = request.args.get('filename', '')
    return f'''
    <h1>Certificado gerado com sucesso!</h1>
    <h4>Confira o diretorio Certificados Gerados onde fica o executavel do programa.</h4><br>
    <h4>Para gerar um novo Certificado basta voltar a página anterior.</h4><br>
    '''

@app.route('/download')
def download():
    # Recupera o nome do arquivo a partir da query string
    filename = request.args.get('filename', '')

    # Caminho completo para o arquivo a ser baixado
    file_path = os.path.join('Certificados_Gerados', filename)

    # Verificar se o arquivo existe
    if not os.path.exists(file_path):
        return "Arquivo não encontrado."

    # Retornar o arquivo para download
    return send_file(file_path, as_attachment=True, download_name=filename)

if __name__ == '__main__':
    app.run(debug=True)
